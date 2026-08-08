"""P1-13b — File intake: PDF / XLSX / CSV / DOCX / image -> text (CPU-only, offline).

Ported from prototype services/intake.py extract_text (proven on txt/pdf/png in demo).
Original file is ALWAYS preserved byte-for-byte under the archive volume and
SHA-256 hashed; AI-visible text is derived separately (v2.0 §3).
"""
import csv
import hashlib
import shutil
import subprocess
from pathlib import Path

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
TEXT_EXTS = {".txt", ".md"}
TABULAR_EXTS = {".xlsx", ".csv"}
SUPPORTED = IMAGE_EXTS | TEXT_EXTS | TABULAR_EXTS | {".pdf", ".docx"}


def sha256_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()


def extract_text(path: Path) -> tuple[str, str]:
    """Return (text, method). Raises ValueError on unsupported/unreadable."""
    ext = path.suffix.lower()
    if ext in TEXT_EXTS:
        return path.read_text(encoding="utf-8", errors="replace"), "text"
    if ext == ".pdf":
        return _pdf_text(path)
    if ext == ".docx":
        return _docx_text(path)
    if ext == ".xlsx":
        return _xlsx_text(path)
    if ext == ".csv":
        return _csv_text(path)
    if ext in IMAGE_EXTS:
        return _ocr_image(path)
    raise ValueError(f"unsupported file type: {ext}")


def _pdf_text(path: Path) -> tuple[str, str]:
    import pdfplumber
    parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
            for table in page.extract_tables() or []:
                for row in table:
                    parts.append(" | ".join(c or "" for c in row))
    text = "\n".join(parts).strip()
    if text:
        return text, "pdfplumber"
    return _ocr_pdf(path)                      # scanned PDF -> rasterize + OCR


def _ocr_pdf(path: Path) -> tuple[str, str]:
    tmp = path.parent / f".{path.stem}_ocr_tmp"
    tmp.mkdir(exist_ok=True)
    try:
        subprocess.run(["pdftoppm", "-r", "200", "-png", str(path), str(tmp / "pg")],
                       check=True, capture_output=True)
        texts = [_ocr_image(p)[0] for p in sorted(tmp.glob("pg-*.png"))]
        text = "\n".join(texts).strip()
        if not text:
            raise ValueError("OCR produced no text (unreadable scan?)")
        return text, "tesseract-pdf"
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def _docx_text(path: Path) -> tuple[str, str]:
    import docx
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts).strip(), "python-docx"


def _xlsx_text(path: Path) -> tuple[str, str]:
    import openpyxl
    wb = openpyxl.load_workbook(str(path), data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"[Sheet: {ws.title}]")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip(), "openpyxl"


def _csv_text(path: Path) -> tuple[str, str]:
    parts = []
    with open(path, newline="", encoding="utf-8", errors="replace") as f:
        for row in csv.reader(f):
            if any(c.strip() for c in row):
                parts.append(" | ".join(row))
    return "\n".join(parts).strip(), "csv"


def _ocr_image(path: Path) -> tuple[str, str]:
    import pytesseract
    from PIL import Image
    text = pytesseract.image_to_string(Image.open(str(path))).strip()
    if not text:
        raise ValueError("OCR produced no text (unreadable image?)")
    return text, "tesseract"
