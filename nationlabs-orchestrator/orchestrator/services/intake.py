"""Intake service (spec §3, §28.1–28.2).

Accepts informal requirements (WhatsApp paste, screenshot, PDF, DOCX, XLSX, image,
voice-note transcript). ALWAYS preserves the original byte-for-byte; AI-visible
text is derived separately. OCR/PDF/DOCX extraction is CPU-only and offline.
"""
from __future__ import annotations

import logging
import shutil
import sqlite3
import subprocess
from datetime import datetime
from pathlib import Path

from ..audit import audit
from ..config import CFG, Config
from ..statemachine import transition

log = logging.getLogger("orchestrator.intake")

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
TEXT_EXTS = {".txt", ".md", ".eml"}


def next_opp_id(conn: sqlite3.Connection) -> str:
    year = datetime.now().year
    row = conn.execute(
        "SELECT opp_id FROM opportunities WHERE opp_id LIKE ? ORDER BY opp_id DESC LIMIT 1",
        (f"NL-OPP-{year}-%",),
    ).fetchone()
    seq = int(row["opp_id"].rsplit("-", 1)[1]) + 1 if row else 1
    return f"NL-OPP-{year}-{seq:04d}"


def extract_text(path: Path) -> tuple[str, str]:
    """Return (text, method). Raises on unsupported/unreadable."""
    ext = path.suffix.lower()
    if ext in TEXT_EXTS:
        return path.read_text(encoding="utf-8", errors="replace"), "text"
    if ext == ".pdf":
        return _pdf_text(path), "pdfplumber"
    if ext == ".docx":
        return _docx_text(path), "python-docx"
    if ext == ".xlsx":
        return _xlsx_text(path), "openpyxl"
    if ext in IMAGE_EXTS:
        return _ocr_text(path), "tesseract"
    raise ValueError(f"unsupported intake type: {ext}")


def _pdf_text(path: Path) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
            for table in page.extract_tables() or []:
                for row in table:
                    parts.append(" | ".join(c or "" for c in row))
    text = "\n".join(parts).strip()
    if not text:  # scanned PDF → OCR fallback
        return _ocr_pdf(path)
    return text


def _ocr_pdf(path: Path) -> str:
    """Rasterize each page via pypdfium2/pdftoppm fallback, then tesseract."""
    tmp = path.parent / f".{path.stem}_ocr_tmp"
    tmp.mkdir(exist_ok=True)
    try:
        subprocess.run(["pdftoppm", "-r", "200", "-png", str(path), str(tmp / "pg")],
                       check=True, capture_output=True)
        texts = [_ocr_text(p) for p in sorted(tmp.glob("pg-*.png"))]
        return "\n".join(texts).strip()
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def _docx_text(path: Path) -> str:
    import docx
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)


def _xlsx_text(path: Path) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(str(path), data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"--- sheet: {ws.title} ---")
        for row in ws.iter_rows(values_only=True):
            if any(v is not None for v in row):
                parts.append(" | ".join("" if v is None else str(v) for v in row))
    return "\n".join(parts)


def _ocr_text(path: Path) -> str:
    try:
        import pytesseract
        from PIL import Image, ImageOps
        img = Image.open(path)
        img = ImageOps.grayscale(img)
        if max(img.size) < 1500:  # upscale small phone screenshots
            img = img.resize((img.width * 2, img.height * 2))
        return pytesseract.image_to_string(img).strip()
    except ImportError:
        # CLI fallback
        r = subprocess.run(["tesseract", str(path), "stdout"],
                           check=True, capture_output=True, text=True)
        return r.stdout.strip()


def create_opportunity(
    conn: sqlite3.Connection,
    *,
    source_channel: str,
    actor: str,
    pasted_text: str | None = None,
    file_path: str | Path | None = None,
    customer_org: str | None = None,
    end_user_org: str | None = None,
    opportunity_owner: str | None = None,
    submission_deadline: str | None = None,
    internal_notes: str | None = None,
    cfg: Config = CFG,
) -> tuple[str, str]:
    """Create opportunity from paste or file. Returns (opp_id, extracted_text).

    Original content is copied unmodified into rfp_archive/<opp_id>/ (spec §3).
    """
    if not pasted_text and not file_path:
        raise ValueError("provide pasted_text or file_path")

    opp_id = next_opp_id(conn)
    archive_dir = cfg.rfp_archive / opp_id
    archive_dir.mkdir(parents=True, exist_ok=True)

    if file_path:
        src = Path(file_path)
        preserved = archive_dir / f"original{src.suffix.lower()}"
        shutil.copy2(src, preserved)                      # never modify the source
        try:
            extracted_text, method = extract_text(preserved)
            (archive_dir / "extracted_text.txt").write_text(extracted_text, encoding="utf-8")
        except Exception as e:
            # Spec §5.3 "Unreadable": intake must survive missing OCR engines,
            # corrupt files, etc. The original is preserved; analysis flags it.
            extracted_text, method = "", f"unreadable:{type(e).__name__}"
            log.warning("text extraction failed for %s: %s", preserved, e)
        source_raw_path = str(preserved)
    else:
        preserved = archive_dir / "original.txt"
        preserved.write_text(pasted_text, encoding="utf-8")
        extracted_text, method = pasted_text, "paste"
        source_raw_path = str(preserved)

    if internal_notes:
        (archive_dir / "internal_notes.txt").write_text(internal_notes, encoding="utf-8")

    with conn:
        conn.execute(
            """INSERT INTO opportunities
               (opp_id, status, customer_org, end_user_org, opportunity_owner,
                source_channel, source_raw_path, submission_deadline)
               VALUES (?,?,?,?,?,?,?,?)""",
            (opp_id, "New Intake", customer_org, end_user_org, opportunity_owner,
             source_channel, source_raw_path, submission_deadline),
        )
        audit(conn, opp_id=opp_id, actor=actor, component="intake",
              action="opportunity_created", new_value=opp_id,
              source=f"{source_channel}/{method}")
        transition(conn, opp_id, "Under Analysis", actor=actor,
                   reason="intake complete, ready for AI extraction")
    log.info("created %s via %s (%s)", opp_id, source_channel, method)
    return opp_id, extracted_text
